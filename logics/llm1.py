
# -----------------------------
# Imports
# -----------------------------
import os, re, requests
from datetime import datetime

import streamlit as st   # Needed for @st.cache_resource
from docx import Document as DocxDocument

from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.embeddings import OpenAIEmbeddings
from langchain.chains import RetrievalQA
from langchain.chat_models import ChatOpenAI
from langchain.agents import Tool, initialize_agent
from langchain.prompts import PromptTemplate

from dotenv import load_dotenv
load_dotenv()


# -----------------------------
# TravelPal Document Loader (RAG)
# -----------------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RAG_PATH = os.path.join(BASE_DIR, "TravelPal RAG document.docx")


@st.cache_resource(show_spinner=False)
def load_travelpal_rag(path=RAG_PATH):
    """
    Loads and embeds the official TravelPal document (as of 10 Nov 2025),
    storing any URLs found in metadata for reference.
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"TravelPal RAG document not found at {path}")

    doc = DocxDocument(path)

    documents = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        urls = re.findall(r'https?://[^\s]+', text)
        documents.append(Document(page_content=text, metadata={"urls": urls}))

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    split_docs = text_splitter.split_documents(documents)

    embeddings = OpenAIEmbeddings(api_key=os.environ.get("OPENAI_API_KEY"))
    vectorstore = FAISS.from_documents(split_docs, embeddings)
    return vectorstore.as_retriever(search_kwargs={"k": 3})


travelpal_retriever = load_travelpal_rag()

# -----------------------------
# LLM Setup
# -----------------------------
llm = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0.4,
    api_key=os.environ.get("OPENAI_API_KEY")
)

prompt = PromptTemplate(
    template=(
        "Answer the question ONLY using the information provided in the context below. "
        "Do NOT use your own knowledge or assume anything.\n\n"
        "Context:\n{context}\n\nQuestion: {question}\nAnswer:"
    ),
    input_variables=["context", "question"]
)


def travelpal_tool_func(query: str):
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        retriever=travelpal_retriever,
        chain_type="stuff",
        chain_type_kwargs={"prompt": prompt}
    )
    answer = qa_chain.run(query)

    docs = travelpal_retriever.get_relevant_documents(query)
    urls = []
    for d in docs:
        if "urls" in d.metadata:
            urls.extend(d.metadata["urls"])
    urls = list(set(urls))

    if urls:
        url_text = "\n".join(f"- {u}" for u in urls)
        return f"{answer}\n\nReference URLs:\n{url_text}"
    return answer


travelpal_tool = Tool(
    name="TravelPal Singapore Policies",
    func=travelpal_tool_func,
    description=(
        "Answer Singapore travel-related questions using TravelPal document extracted from the following URLs:" \
        "Travel Tips from MFA:https://www.mfa.gov.sg/Consular-Services/Singapore-Citizens/Travel-Tips" \
        "Need help from the MHA while overseas: https://www.mfa.gov.sg/Consular-Services/Singapore-Citizens/I-Need-Help-Overseas"\
        "Prohibited, Controlled, and Dutiable Goods – Entering Singapore: https://www.ica.gov.sg/enter-transit-depart/entering-singapore/what-you-can-bring/prohibited-controlled-dutiable-goods"\
        "Advice from ICA for Singapore Citizens Travelling Abroad/ Returning to Singapore: https://www.ica.gov.sg/enter-depart/for-singapore-citizens/advice-for-travelling-abroad "\
        "APEC Business Travel Card:https://www.ica.gov.sg/enter-depart/for-singapore-citizens/apec-business-travel-card "
        "Always include reference URLs."
    ),
)

# -----------------------------
# MFA Country Page Tool
# -----------------------------
def fetch_mfa_page(country: str):
    country_map = {
        "Japan": "https://www.mfa.gov.sg/countries-regions/j/japan/travel-page",
        "China": "https://www.mfa.gov.sg/countries-regions/c/china/travel-page",
        "Thailand": "https://www.mfa.gov.sg/countries-regions/t/thailand/travel-page",
        "Australia": "https://www.mfa.gov.sg/countries-regions/a/australia/travel-page",
    }
    url = country_map.get(country.title())
    if not url:
        return f"I couldn’t find a dedicated MFA advisory page for {country}."
    return f"See MFA advisory for {country}: {url}"


mfa_tool = Tool(
    name="MFA Country Advisory Tool",
    func=lambda q: fetch_mfa_page(q.split()[-1]),
    description="Country-specific MFA travel advisories.",
)

# -----------------------------
# Weather Tool
# -----------------------------
def get_weather(city: str, month: int):
    """
    Retrieves average temperature for a given city and month.
    """
    # Use geocoding API to get lat/lon
    url = f"https://geocoding-api.open-meteo.com/v1/search?name={city}"
    loc = requests.get(url).json()
    if "results" not in loc:
        return f"Sorry, I couldn’t find {city}."
    lat, lon = loc["results"][0]["latitude"], loc["results"][0]["longitude"]

    weather_url = f"https://climate-api.open-meteo.com/v1/climate?latitude={lat}&longitude={lon}&month={month}"
    wdata = requests.get(weather_url).json()
    temp = wdata.get("data", {}).get("temperature_2m_mean")

    if temp is not None:
        return f"The average temperature in {city.title()} in month {month} is around {temp}°C."
    return "Weather data unavailable."


def weather_tool_func(query: str):
    match = re.search(r'in ([A-Za-z\s]+)', query)
    city = match.group(1).strip() if match else query.strip()
    month = datetime.now().month
    return get_weather(city, month)


weather_tool = Tool(
    name="Weather Helper",
    func=weather_tool_func,
    description="Provides average monthly temperature for a given city.",
)

# -----------------------------
# Assemble Tools & Initialize Agent
# -----------------------------
tools = [travelpal_tool, mfa_tool, weather_tool]

agent = initialize_agent(
    tools=tools,
    llm=llm,
    agent_type="zero-shot-react-description",
    verbose=False,
)

__all__ = ["agent"]
