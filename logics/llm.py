# -----------------------------
# Imports
# -----------------------------
import os, re, requests
from datetime import datetime
from bs4 import BeautifulSoup

import streamlit as st
from docx import Document as DocxDocument

from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.embeddings import OpenAIEmbeddings
from langchain.chains import RetrievalQA
from langchain.chat_models import ChatOpenAI
from langchain.agents import Tool, initialize_agent
from langchain.prompts import PromptTemplate

from dotenv import load_dotenv
load_dotenv()

# -----------------------------
# Base Directories & Paths
# -----------------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RAG_PATH = os.path.join(BASE_DIR, "TravelPal RAG document.docx")

# -----------------------------
# Helper: Extract Country
# -----------------------------
import spacy
nlp = spacy.load("en_core_web_sm")

def extract_country(query: str):
    doc = nlp(query)
    countries = [ent.text for ent in doc.ents if ent.label_ == "GPE"]
    return countries[0] if countries else None

# -----------------------------
# TravelPal RAG Loader
# -----------------------------
@st.cache_resource(show_spinner=False)
def load_travelpal_rag(path=RAG_PATH):
    if not os.path.exists(path):
        raise FileNotFoundError(f"TravelPal RAG document not found at {path}")
    doc = DocxDocument(path)

    documents = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        urls = re.findall(r'https?://[^\s]+', text)
        documents.append(Document(page_content=text, metadata={"urls": urls}))

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    split_docs = text_splitter.split_documents(documents)

    embeddings = OpenAIEmbeddings(api_key=os.environ.get("OPENAI_API_KEY"))
    vectorstore = FAISS.from_documents(split_docs, embeddings)
    return vectorstore.as_retriever(search_kwargs={"k": 3})

travelpal_retriever = load_travelpal_rag()

# -----------------------------
# LLM Setup
# -----------------------------
llm = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0.4,
    api_key=os.environ.get("OPENAI_API_KEY")
)

prompt = PromptTemplate(
    template=(
        "Answer the question ONLY using the information provided in the context below. "
        "Do NOT use your own knowledge or assume anything.\n\n"
        "Context:\n{context}\n\nQuestion: {question}\nAnswer:"
    ),
    input_variables=["context", "question"]
)

# -----------------------------
# TravelPal Tool
# -----------------------------
def travelpal_tool_func(query: str):
    # Run RetrievalQA
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        retriever=travelpal_retriever,
        chain_type="stuff",
        chain_type_kwargs={"prompt": prompt}
    )
    answer = qa_chain.run(query)

    # Collect relevant URLs from the retrieved documents
    docs = travelpal_retriever.get_relevant_documents(query)
    urls = []
    for d in docs:
        if "urls" in d.metadata:
            urls.extend(d.metadata["urls"])
    urls = list(set(urls))  # Remove duplicates

    # Make URLs clickable in Markdown
    if urls:
        url_text = "\n".join(f"[{u}]({u})" for u in urls)
        answer += f"\n\n**Reference URLs:**\n{url_text}"

    return answer


travelpal_tool = Tool(
    name="TravelPal Singapore Policies",
    func=travelpal_tool_func,
    description=(
        "Use this tool to answer Singapore travel-related questions. "
        "You must retrieve information from the following official sources:\n"
        "- Travel Tips from MFA (before travelling and while already overseas): https://www.mfa.gov.sg/Consular-Services/Singapore-Citizens/Travel-Tips\n"
        "- MFA Assistance for Singaporeans Overseas (how to get help abroad): https://www.mfa.gov.sg/Consular-Services/Singapore-Citizens/I-Need-Help-Overseas\n"
        "- ICA Guidelines on Prohibited, Controlled, and Dutiable Goods when entering Singapore: https://www.ica.gov.sg/enter-transit-depart/entering-singapore/what-you-can-bring/prohibited-controlled-dutiable-goods\n"
        "- ICA advice for Singapore citizens travelling abroad or returning to Singapore: https://www.ica.gov.sg/enter-depart/for-singapore-citizens/advice-for-travelling-abroad\n"
        "- APEC Business Travel Card information: https://www.ica.gov.sg/enter-depart/for-singapore-citizens/apec-business-travel-card\n\n"
        "Instructions for the tool output:\n"
        "1. Summarise information from the relevant sources to answer the user's question.\n"
        "2. Include the specific reference URL(s) used in the answer.\n"
        "3. Do NOT tell the user to refer to TravelPal policies.\n"
        "4. Provide a clear, concise, and factual response.\n"
    ),
)

# -----------------------------
# MFA Tool
# -----------------------------

MFA_COUNTRY_MAP = {
    "Afghanistan": "https://www.mfa.gov.sg/countries-regions/a/afghanistan/travel-page",
    "Albania": "https://www.mfa.gov.sg/countries-regions/a/albania/travel-page",
    "Algeria": "https://www.mfa.gov.sg/countries-regions/a/algeria/travel-page",
    "Angola": "https://www.mfa.gov.sg/countries-regions/a/angola/travel-page",
    "Antigua and Barbuda": "https://www.mfa.gov.sg/countries-regions/a/antigua-and-barbuda/travel-page",
    "Argentina": "https://www.mfa.gov.sg/countries-regions/a/argentina/travel-page",
    "Armenia": "https://www.mfa.gov.sg/countries-regions/a/armenia/travel-page",
    "Australia": "https://www.mfa.gov.sg/countries-regions/a/australia/travel-page",
    "Austria": "https://www.mfa.gov.sg/countries-regions/a/austria/travel-page",
    "Azerbaijan": "https://www.mfa.gov.sg/countries-regions/a/azerbaijan/travel-page",
    "Bahamas": "https://www.mfa.gov.sg/countries-regions/b/bahamas/travel-page",
    "Bahrain": "https://www.mfa.gov.sg/countries-regions/b/bahrain/travel-page",
    "Bangladesh": "https://www.mfa.gov.sg/countries-regions/b/bangladesh/travel-page",
    "Barbados": "https://www.mfa.gov.sg/countries-regions/b/barbados/travel-page",
    "Belarus": "https://www.mfa.gov.sg/countries-regions/b/belarus/travel-page",
    "Belgium": "https://www.mfa.gov.sg/countries-regions/b/belgium/travel-page",
    "Belize": "https://www.mfa.gov.sg/countries-regions/b/belize/travel-page",
    "Benin": "https://www.mfa.gov.sg/countries-regions/b/benin/travel-page",
    "Bhutan": "https://www.mfa.gov.sg/countries-regions/b/bhutan/travel-page",
    "Bolivia": "https://www.mfa.gov.sg/countries-regions/b/bolivia-plurinational-state-of/travel-page",
    "Bosnia and Herzegovina": "https://www.mfa.gov.sg/countries-regions/b/bosnia-and-herzegovina/travel-page",
    "Botswana": "https://www.mfa.gov.sg/countries-regions/b/botswana/travel-page",
    "Brazil": "https://www.mfa.gov.sg/countries-regions/b/brazil/travel-page",
    "Brunei": "https://www.mfa.gov.sg/countries-regions/b/brunei-darussalam/travel-page",
    "Bulgaria": "https://www.mfa.gov.sg/countries-regions/b/bulgaria/travel-page",
    "Burkina Faso": "https://www.mfa.gov.sg/countries-regions/b/burkina-faso/travel-page",
    "Cabo Verde": "https://www.mfa.gov.sg/countries-regions/c/cabo-verde/travel-page",
    "Cambodia": "https://www.mfa.gov.sg/countries-regions/c/cambodia/travel-page",
    "Cameroon": "https://www.mfa.gov.sg/countries-regions/c/cameroon/travel-page",
    "Canada": "https://www.mfa.gov.sg/countries-regions/c/canada/travel-page",
    "Chad": "https://www.mfa.gov.sg/countries-regions/c/chad/travel-page",
    "Chile": "https://www.mfa.gov.sg/countries-regions/c/chile/travel-page",
    "China": "https://www.mfa.gov.sg/countries-regions/c/china/travel-page",
    "Colombia": "https://www.mfa.gov.sg/countries-regions/c/colombia/travel-page",
    "Comoros": "https://www.mfa.gov.sg/countries-regions/c/comoros/travel-page",
    "Congo": "https://www.mfa.gov.sg/countries-regions/c/congo/travel-page",
    "Democratic Republic of Congo": "https://www.mfa.gov.sg/countries-regions/d/democratic-republic-of-congo/travel-page",
    "Cook Islands": "https://www.mfa.gov.sg/countries-regions/c/cook-islands/travel-page",
    "Costa Rica": "https://www.mfa.gov.sg/countries-regions/c/costa-rica/travel-page",
    "Cote d Ivoire": "https://www.mfa.gov.sg/countries-regions/c/cote-d-ivoire/travel-page",
    "Croatia": "https://www.mfa.gov.sg/countries-regions/c/croatia/travel-page",
    "Cuba": "https://www.mfa.gov.sg/countries-regions/c/cuba/travel-page",
    "Czech Republic": "https://www.mfa.gov.sg/countries-regions/c/czech-republic/travel-page",
    "Denmark": "https://www.mfa.gov.sg/countries-regions/d/denmark/travel-page",
    "Djibouti": "https://www.mfa.gov.sg/countries-regions/d/djibouti/travel-page",
    "Dominica": "https://www.mfa.gov.sg/countries-regions/d/dominica/travel-page",
    "Dominican Republic": "https://www.mfa.gov.sg/countries-regions/d/dominican-republic/travel-page",
    "Ecuador": "https://www.mfa.gov.sg/countries-regions/e/ecuador/travel-page",
    "Egypt": "https://www.mfa.gov.sg/countries-regions/e/egypt/travel-page",
    "El Salvador": "https://www.mfa.gov.sg/countries-regions/e/el-salvador/travel-page",
    "Estonia": "https://www.mfa.gov.sg/countries-regions/e/estonia/travel-page",
    "Eswatini": "https://www.mfa.gov.sg/countries-regions/e/eswatini/travel-page",
    "Ethiopia": "https://www.mfa.gov.sg/countries-regions/e/ethiopia/travel-page",
    "Federated States of Micronesia": "https://www.mfa.gov.sg/countries-regions/f/federated-states-of-micronesia/travel-page",
    "Fiji": "https://www.mfa.gov.sg/countries-regions/f/fiji/travel-page",
    "Finland": "https://www.mfa.gov.sg/countries-regions/f/finland/travel-page",
    "France": "https://www.mfa.gov.sg/countries-regions/f/france/travel-page",
    "Gabon": "https://www.mfa.gov.sg/countries-regions/g/gabon/travel-page",
    "Gambia": "https://www.mfa.gov.sg/countries-regions/g/gambia/travel-page",
    "Georgia": "https://www.mfa.gov.sg/countries-regions/g/georgia/travel-page",
    "Germany": "https://www.mfa.gov.sg/countries-regions/g/germany/travel-page",
    "Ghana": "https://www.mfa.gov.sg/countries-regions/g/ghana/travel-page",
    "Greece": "https://www.mfa.gov.sg/countries-regions/g/greece/travel-page",
    "Grenada": "https://www.mfa.gov.sg/countries-regions/g/grenada/travel-page",
    "Guatemala": "https://www.mfa.gov.sg/countries-regions/g/guatemala/travel-page",
    "Republic of Guinea": "https://www.mfa.gov.sg/countries-regions/g/republic-of-guinea/travel-page",
    "Guinea-Bissau": "https://www.mfa.gov.sg/countries-regions/g/guinea-bissau/travel-page",
    "Guyana": "https://www.mfa.gov.sg/countries-regions/g/guyana/travel-page",
    "Haiti": "https://www.mfa.gov.sg/countries-regions/h/haiti/travel-page",
    "Honduras": "https://www.mfa.gov.sg/countries-regions/h/honduras/travel-page",
    "Hong Kong": "https://www.mfa.gov.sg/countries-regions/h/hong-kong/travel-page",
    "Hungary": "https://www.mfa.gov.sg/countries-regions/h/hungary/travel-page",
    "Iceland": "https://www.mfa.gov.sg/countries-regions/i/iceland/travel-page",
    "India": "https://www.mfa.gov.sg/countries-regions/i/india/travel-page",
    "Indonesia": "https://www.mfa.gov.sg/countries-regions/i/indonesia/travel-page",
    "Iran": "https://www.mfa.gov.sg/countries-regions/i/iran-islamic-republic-of/travel-page",
    "Iraq": "https://www.mfa.gov.sg/countries-regions/i/iraq/travel-page",
    "Ireland": "https://www.mfa.gov.sg/countries-regions/i/ireland/travel-page",
    "Israel": "https://www.mfa.gov.sg/countries-regions/i/israel/travel-page",
    "Italy": "https://www.mfa.gov.sg/countries-regions/i/italy/travel-page",
    "Jamaica": "https://www.mfa.gov.sg/countries-regions/j/jamaica/travel-page",
    "Japan": "https://www.mfa.gov.sg/countries-regions/j/japan/travel-page",
    "Jordan": "https://www.mfa.gov.sg/countries-regions/j/jordan/travel-page",
    "Kazakhstan": "https://www.mfa.gov.sg/countries-regions/k/kazakhstan/travel-page",
    "Kenya": "https://www.mfa.gov.sg/countries-regions/k/kenya/travel-page",
    "Kiribati": "https://www.mfa.gov.sg/countries-regions/k/kiribati/travel-page",
    "North Korea": "https://www.mfa.gov.sg/countries-regions/k/korea-democratic-peoples-republic-of/travel-page",
    "South Korea": "https://www.mfa.gov.sg/countries-regions/k/korea-republic-of/travel-page",
    "Kuwait": "https://www.mfa.gov.sg/countries-regions/k/kuwait/travel-page",
    "Kyrgyz Republic": "https://www.mfa.gov.sg/countries-regions/k/kyrgyz-republic/travel-page",
    "Laos": "https://www.mfa.gov.sg/countries-regions/l/lao-peoples-democratic-republic/travel-page",
    "Latvia":"https://www.mfa.gov.sg/countries-regions/l/latvia/travel-page",
    "Lebanon": "https://www.mfa.gov.sg/countries-regions/l/lebanon/travel-page",
    "Lesotho": "https://www.mfa.gov.sg/countries-regions/l/lesotho/travel-page",
    "Liberia": "https://www.mfa.gov.sg/countries-regions/l/liberia/travel-page",
    "Libya": "https://www.mfa.gov.sg/countries-regions/l/libya/travel-page",
    "Liechtenstein": "https://www.mfa.gov.sg/countries-regions/l/liechtenstein/travel-page",
    "Lithuania": "https://www.mfa.gov.sg/countries-regions/l/lithuania/travel-page",
    "Luxembourg": "https://www.mfa.gov.sg/countries-regions/l/luxembourg/travel-page",
    "Macao": "https://www.mfa.gov.sg/countries-regions/m/macao/travel-page",
    "Madagascar": "https://www.mfa.gov.sg/countries-regions/m/madagascar/travel-page",
    "Malawi": "https://www.mfa.gov.sg/countries-regions/m/malawi/travel-page",
    "Malaysia": "https://www.mfa.gov.sg/countries-regions/m/malaysia/travel-page",
    "Maldives": "https://www.mfa.gov.sg/countries-regions/m/maldives/travel-page",
    "Mali": "https://www.mfa.gov.sg/countries-regions/m/mali/travel-page",
    "Marshall Islands": "https://www.mfa.gov.sg/countries-regions/m/marshall-islands/travel-page",
    "Mauritania": "https://www.mfa.gov.sg/countries-regions/m/mauritania/travel-page",
    "Mauritius": "https://www.mfa.gov.sg/countries-regions/m/mauritius/travel-page",
    "Mexico": "https://www.mfa.gov.sg/countries-regions/m/mexico/travel-page",
    "Moldova": "https://www.mfa.gov.sg/countries-regions/m/moldova/travel-page",
    "Mongolia": "https://www.mfa.gov.sg/countries-regions/m/mongolia/travel-page",
    "Montenegro": "https://www.mfa.gov.sg/countries-regions/m/montenegro/travel-page",
    "Morocco": "https://www.mfa.gov.sg/countries-regions/m/morocco/travel-page",
    "Mozambique": "https://www.mfa.gov.sg/countries-regions/m/mozambique/travel-page",
    "Myanmar": "https://www.mfa.gov.sg/countries-regions/m/myanmar/travel-page",
    "Namibia": "https://www.mfa.gov.sg/countries-regions/n/namibia/travel-page",
    "Nauru": "https://www.mfa.gov.sg/countries-regions/n/nauru/travel-page",
    "Nepal": "https://www.mfa.gov.sg/countries-regions/n/nepal/travel-page",
    "Netherlands": "https://www.mfa.gov.sg/countries-regions/n/netherlands/travel-page",
    "New zealand": "https://www.mfa.gov.sg/countries-regions/n/new-zealand/travel-page",
    "Nicaragua": "https://www.mfa.gov.sg/countries-regions/n/nicaragua/travel-page",
    "Niger": "https://www.mfa.gov.sg/countries-regions/n/niger/travel-page",
    "Nigeria": "https://www.mfa.gov.sg/countries-regions/n/nigeria/travel-page",
    "Niue": "https://www.mfa.gov.sg/countries-regions/n/niue/travel-page",
    "North Macedonia": "https://www.mfa.gov.sg/countries-regions/n/north-macedonia/travel-page",
    "Norway": "https://www.mfa.gov.sg/countries-regions/n/norway/travel-page",
    "Oman": "https://www.mfa.gov.sg/countries-regions/o/oman/travel-page",
    "Pakistan": "https://www.mfa.gov.sg/countries-regions/p/pakistan/travel-page",
    "Palau": "https://www.mfa.gov.sg/countries-regions/p/palau/travel-page",
    "Palestinian Territories": "https://www.mfa.gov.sg/countries-regions/p/palestinian-territories/travel-page",
    "Panama": "https://www.mfa.gov.sg/countries-regions/p/panama/travel-page",
    "Papua New Guinea": "https://www.mfa.gov.sg/countries-regions/p/papua-new-guinea/travel-page",
    "Paraguay": "https://www.mfa.gov.sg/countries-regions/p/paraguay/travel-page",
    "Peru": "https://www.mfa.gov.sg/countries-regions/p/peru/travel-page",
    "Philippines": "https://www.mfa.gov.sg/countries-regions/p/philippines/travel-page",
    "Poland": "https://www.mfa.gov.sg/countries-regions/p/poland/travel-page",
    "Portugal": "https://www.mfa.gov.sg/countries-regions/p/portugal/travel-page",
    "Qatar": "https://www.mfa.gov.sg/countries-regions/q/qatar/travel-page",
    "Romania": "https://www.mfa.gov.sg/countries-regions/r/romania/travel-page",
    "Russia": "https://www.mfa.gov.sg/countries-regions/r/russian-federation/travel-page",
    "Rwanda": "https://www.mfa.gov.sg/countries-regions/r/rwanda/travel-page",
    "Saint Kitts and Nevis": "https://www.mfa.gov.sg/countries-regions/s/saint-kitts-and-nevis/travel-page",
    "Saint Lucia": "https://www.mfa.gov.sg/countries-regions/s/saint-lucia/travel-page",
    "Saint Vincent and the Grenadines": "https://www.mfa.gov.sg/countries-regions/s/saint-vincent-and-the-grenadines/travel-page",
    "Samoa": "https://www.mfa.gov.sg/countries-regions/s/samoa/travel-page",
    "Saudi Arabia": "https://www.mfa.gov.sg/countries-regions/s/saudi-arabia/travel-page",
    "Senegal": "https://www.mfa.gov.sg/countries-regions/s/senegal/travel-page",
    "Serbia": "https://www.mfa.gov.sg/countries-regions/s/serbia/travel-page",
    "Seychelles": "https://www.mfa.gov.sg/countries-regions/s/seychelles/travel-page",
    "Sierra Leone": "https://www.mfa.gov.sg/countries-regions/s/sierra-leone/travel-page",
    "Slovakia": "https://www.mfa.gov.sg/countries-regions/s/slovakia/travel-page",
    "Slovenia": "https://www.mfa.gov.sg/countries-regions/s/slovenia/travel-page",
    "Solomon Islands": "https://www.mfa.gov.sg/countries-regions/s/solomon-islands/travel-page",
    "Somalia": "https://www.mfa.gov.sg/countries-regions/s/somalia/travel-page",
    "South Africa": "https://www.mfa.gov.sg/countries-regions/s/south-africa/travel-page",
    "Spain": "https://www.mfa.gov.sg/countries-regions/s/spain/travel-page",
    "Sri Lanka": "https://www.mfa.gov.sg/countries-regions/s/sri-lanka/travel-page",
    "Suriname": "https://www.mfa.gov.sg/countries-regions/s/suriname/travel-page",
    "Sweden": "https://www.mfa.gov.sg/countries-regions/s/sweden/travel-page",
    "Switzerland": "https://www.mfa.gov.sg/countries-regions/s/switzerland/travel-page",
    "Syria": "https://www.mfa.gov.sg/countries-regions/s/syrian-arab-republic/travel-page",
    "Taiwan": "https://www.mfa.gov.sg/countries-regions/t/taiwan/travel-page",
    "Tajikistan": "https://www.mfa.gov.sg/countries-regions/t/tajikistan/travel-page",
    "Tanzania": "https://www.mfa.gov.sg/countries-regions/t/tanzania-united-republic-of/travel-page",
    "Thailand": "https://www.mfa.gov.sg/countries-regions/t/thailand/travel-page",
    "Timor-Leste": "https://www.mfa.gov.sg/countries-regions/t/timor-leste/travel-page",
    "Togo": "https://www.mfa.gov.sg/countries-regions/t/togo/travel-page",
    "Tonga": "https://www.mfa.gov.sg/countries-regions/t/tonga/travel-page",
    "Trinidad and Tobago": "https://www.mfa.gov.sg/countries-regions/t/trinidad-and-tobago/travel-page",
    "Tunisia": "https://www.mfa.gov.sg/countries-regions/t/tunisia/travel-page",
    "Turkiye": "https://www.mfa.gov.sg/countries-regions/t/turkiye/travel-page",
    "Turkmenistan": "https://www.mfa.gov.sg/countries-regions/t/turkmenistan/travel-page",
    "Tuvalu": "https://www.mfa.gov.sg/countries-regions/t/tuvalu/travel-page",
    "Uganda": "https://www.mfa.gov.sg/countries-regions/u/uganda/travel-page",
    "Ukraine": "https://www.mfa.gov.sg/countries-regions/u/ukraine/travel-page",
    "United Arab Emirates": "https://www.mfa.gov.sg/countries-regions/u/united-arab-emirates/travel-page",
    "United Kingdom": "https://www.mfa.gov.sg/countries-regions/u/united-kingdom/travel-page",
    "United States": "https://www.mfa.gov.sg/countries-regions/u/united-states/travel-page",
    "Uruguay": "https://www.mfa.gov.sg/countries-regions/u/uruguay/travel-page",
    "Uzbekistan": "https://www.mfa.gov.sg/countries-regions/u/uzbekistan/travel-page",
    "Vanuatu": "https://www.mfa.gov.sg/countries-regions/v/vanuatu/travel-page",
    "Venezuela": "https://www.mfa.gov.sg/countries-regions/v/venezuela-bolivarian-republic-of/travel-page",
    "Vietnam": "https://www.mfa.gov.sg/countries-regions/v/viet-nam/travel-page",
    "Yemen": "https://www.mfa.gov.sg/countries-regions/y/yemen/travel-page",
    "Zambia": "https://www.mfa.gov.sg/countries-regions/z/zambia/travel-page",
    "Zimbabwe": "https://www.mfa.gov.sg/countries-regions/z/zimbabwe/travel-page"
}

def mfa_tool_func(query: str):
    country = extract_country(query)
    if not country or country.title() not in MFA_COUNTRY_MAP:
        return "I couldn’t detect a valid country for the MFA advisory."
    
    url = MFA_COUNTRY_MAP[country.title()]
    try:
        r = requests.get(url, timeout=5)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")
        title_text = soup.title.string.strip() if soup.title else f"MFA Travel Advisory for {country}"
    except:
        title_text = f"MFA Travel Advisory for {country}"

    return f"{title_text}: [{url}]({url})"

mfa_tool = Tool(
    name="MFA Country Advisory Tool",
    func=mfa_tool_func,
    description=("Use this tool to answer country-specific travel questions based on official MFA "
        "travel advisory content. The tool contains information for each country in the following categories:\n\n"
        
        "1. Country travel advisories and alerts\n"
        "2. Entry and exit requirements\n"
        "3. Safety and security information\n"
        "4. Local laws and regulations\n"
        "5. General travel advice and precautions\n"
        "6. Local emergency contact numbers\n"
        "7. Mission contact details\n\n"

        #  "This tool must NOT be used for weather-related queries. "
        # "For any question involving weather, temperatures, forecasts, seasons, or climate conditions, "
        # "you must activate `weather_tool_func` instead.\n\n"
        
        "When responding to the user:\n"
        "- Retrieve relevant information only from the provided MFA country advisory content.\n"
        "- Summarise the information clearly and accurately.\n"
        "- Include the specific clickable URL(s) used as references.\n"
        "- Ensure the answer is user-friendly, factual, and concise.\n"
    )
)

# -----------------------------
# Weather Tool
# -----------------------------
def get_weather(city: str, month: int):
    url = f"https://geocoding-api.open-meteo.com/v1/search?name={city}"
    loc = requests.get(url).json()
    if "results" not in loc:
        return f"Sorry, I couldn’t find {city}."
    lat, lon = loc["results"][0]["latitude"], loc["results"][0]["longitude"]

    weather_url = f"https://climate-api.open-meteo.com/v1/climate?latitude={lat}&longitude={lon}&month={month}"
    wdata = requests.get(weather_url).json()
    temp = wdata.get("data", {}).get("temperature_2m_mean")

    if temp is not None:
        return f"The average temperature in {city.title()} in month {month} is around {temp}°C."
    return "Weather data unavailable."

def weather_tool_func(query: str):
    match = re.search(r'in ([A-Za-z\s]+)', query)
    city = match.group(1).strip() if match else query.strip()
    month = datetime.now().month
    return get_weather(city, month)

weather_tool = Tool(
    name="Weather Helper",
    func=weather_tool_func,
    description="Provides average monthly temperature for a given city.",
)

# -----------------------------
# Assemble Tools & Initialize Agent
# -----------------------------
tools = [travelpal_tool, mfa_tool, weather_tool]

agent = initialize_agent(
    tools=tools,
    llm=llm,
    agent_type="zero-shot-react-description",
    verbose=False,
    handle_parsing_errors= True
)

# llm.py
__all__ = ["agent", "travelpal_tool_func", "mfa_tool", "weather_tool_func"]



